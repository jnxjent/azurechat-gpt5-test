"""
edit_word.py  –  python-docx ベースの Word 編集スクリプト
CLI: python edit_word.py --input <path> --output <path> --plan <json_path>
stdout: JSON {"changedParagraphs": N, "totalParagraphs": N}

plan.trackChanges=true の場合、変更をWordの変更履歴（Track Changes）として記録し、
各変更箇所にWordコメントを挿入する。
  - replaceText: 削除テキストを <w:del>、挿入テキストを <w:ins> でマーク、変更箇所にコメント挿入
  - formatRuns: 元の書式を <w:rPrChange> で保持、変更段落にコメント挿入
  - addParagraphs: 追加段落全体を <w:ins> でマーク
"""

import argparse
import copy
import json
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from lxml import etree

TRACK_AUTHOR = "AzureChat"
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)


def parse_hex_color(value: str | None):
    """6桁 RRGGBB → RGBColor。無効な値は None を返す。"""
    if not value:
        return None
    normalized = str(value).replace("#", "").strip().upper()
    if len(normalized) == 6 and all(ch in "0123456789ABCDEF" for ch in normalized):
        r = int(normalized[0:2], 16)
        g = int(normalized[2:4], 16)
        b = int(normalized[4:6], 16)
        return RGBColor(r, g, b)
    return None


def iter_all_paragraphs(doc):
    """段落とテーブルセル内の段落をすべてイテレートする。"""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


# ─── Comments Manager ────────────────────────────────────────────────────────

class CommentsManager:
    """Manages the word/comments.xml part for inserting Word comments."""

    def __init__(self, doc, rev: list):
        self.doc = doc
        self.rev = rev  # shared [int] counter — all w:id values must be unique
        self._part = None
        self._root = None

    def _ensure_part(self):
        if self._part is not None:
            return
        try:
            self._part = self.doc.part.part_related_by(COMMENTS_REL_TYPE)
            self._root = etree.fromstring(self._part.blob)
        except KeyError:
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI

            self._root = OxmlElement("w:comments")
            xml_bytes = etree.tostring(
                self._root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone=True,
            )
            self._part = Part(
                PackURI("/word/comments.xml"),
                COMMENTS_CONTENT_TYPE,
                xml_bytes,
                self.doc.part.package,
            )
            self.doc.part.relate_to(self._part, COMMENTS_REL_TYPE)

    def add_comment(self, text: str, author: str, date: str) -> int:
        """Add a comment entry to comments.xml. Returns the comment w:id used."""
        self._ensure_part()
        comment_id = self.rev[0]
        self.rev[0] += 1

        comment = OxmlElement("w:comment")
        comment.set(qn("w:id"), str(comment_id))
        comment.set(qn("w:author"), author)
        comment.set(qn("w:date"), date)
        comment.set(qn("w:initials"), "AC")

        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "CommentText")
        pPr.append(pStyle)
        p.append(pPr)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "CommentText")
        rPr.append(rStyle)
        r.append(rPr)

        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        comment.append(p)
        self._root.append(comment)
        return comment_id

    def flush(self):
        """Serialize updated comments XML back to the part blob."""
        if self._part is None:
            return
        self._part._blob = etree.tostring(
            self._root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )


# ─── 変更履歴なし（既存実装） ───────────────────────────────────────────────

def apply_replace_text(doc, replacements: list) -> int:
    """replaceText を文書全体（段落・テーブル）に適用。変更した段落数を返す。"""
    changed = 0
    for para in iter_all_paragraphs(doc):
        para_changed = False
        for rep in replacements:
            find = rep.get("find", "")
            replace = rep.get("replace", "")
            if not find:
                continue
            single_run_found = False
            for run in para.runs:
                if find in run.text:
                    run.text = run.text.replace(find, replace)
                    para_changed = True
                    single_run_found = True
            # Cross-run fallback: find が複数 run にまたがる場合
            if not single_run_found and find in para.text:
                if _replace_text_cross_run(para, find, replace):
                    para_changed = True
        if para_changed:
            changed += 1
    return changed


def apply_format_runs(doc, formats: list) -> int:
    """formatRuns を段落に適用。matchText 省略時は全段落に適用。変更した段落数を返す。"""
    changed = 0
    for fmt in formats:
        match_text = str(fmt.get("matchText") or "").strip()
        apply_all = not match_text
        bold = fmt.get("bold")
        italic = fmt.get("italic")
        font_size = fmt.get("fontSize")
        font_color = parse_hex_color(fmt.get("fontColor"))
        font_face = str(fmt.get("fontFace") or "").strip() or None

        if bold is None and italic is None and font_size is None and font_color is None and font_face is None:
            continue

        for para in iter_all_paragraphs(doc):
            if not apply_all and match_text not in para.text:
                continue
            runs = para.runs
            if not runs and para.text:
                run = para.add_run(para.text)
                runs = [run]
            for run in runs:
                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if font_size is not None:
                    run.font.size = Pt(float(font_size))
                if font_color is not None:
                    run.font.color.rgb = font_color
                if font_face is not None:
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), font_face)
                    rFonts.set(qn("w:hAnsi"), font_face)
                    rFonts.set(qn("w:eastAsia"), font_face)
                    rFonts.set(qn("w:cs"), font_face)
            changed += 1

    return changed


def apply_add_paragraphs(doc, paragraphs: list) -> int:
    """addParagraphs をドキュメント末尾に追加。追加した段落数を返す。"""
    added = 0
    for para_def in paragraphs:
        text = str(para_def.get("text") or "").strip()
        if not text:
            continue
        style = str(para_def.get("style") or "Normal").strip()
        bold = para_def.get("bold")
        italic = para_def.get("italic")
        font_size = para_def.get("fontSize")
        font_color = parse_hex_color(para_def.get("fontColor"))

        try:
            para = doc.add_paragraph(text, style=style)
        except Exception:
            para = doc.add_paragraph(text)

        for run in para.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size is not None:
                run.font.size = Pt(float(font_size))
            if font_color is not None:
                run.font.color.rgb = font_color
        added += 1

    return added


# ─── 変更履歴あり（Track Changes） ─────────────────────────────────────────

def _rPr_copy(run_el):
    """run の rPr 要素の deepcopy を返す。存在しない場合は空の w:rPr を返す。"""
    rPr = run_el.find(qn("w:rPr"))
    return copy.deepcopy(rPr) if rPr is not None else OxmlElement("w:rPr")


def _make_del_run(rPr_elem, text: str):
    """<w:del> 内に入る <w:r><w:delText>text</w:delText></w:r> を生成する。"""
    r = OxmlElement("w:r")
    r.append(copy.deepcopy(rPr_elem))
    del_text = OxmlElement("w:delText")
    del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    del_text.text = text
    r.append(del_text)
    return r


def _make_ins_run(rPr_elem, text: str):
    """<w:ins> 内に入る <w:r><w:t>text</w:t></w:r> を生成する。"""
    r = OxmlElement("w:r")
    r.append(copy.deepcopy(rPr_elem))
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _make_plain_run(rPr_elem, text: str):
    """変更なしの通常 <w:r> を生成する。"""
    r = OxmlElement("w:r")
    r.append(copy.deepcopy(rPr_elem))
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap_del(rev_id: int, author: str, date: str, inner_run):
    """inner_run を <w:del> で包んで返す。"""
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    el.append(inner_run)
    return el


def _wrap_ins(rev_id: int, author: str, date: str, inner_run):
    """inner_run を <w:ins> で包んで返す。"""
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    el.append(inner_run)
    return el


def _make_comment_ref_run(comment_id: int) -> Any:
    """コメント参照マーク <w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:commentReference/></w:r>"""
    r_ref = OxmlElement("w:r")
    rPr_ref = OxmlElement("w:rPr")
    rStyle_ref = OxmlElement("w:rStyle")
    rStyle_ref.set(qn("w:val"), "CommentReference")
    rPr_ref.append(rStyle_ref)
    r_ref.append(rPr_ref)
    ref_el = OxmlElement("w:commentReference")
    ref_el.set(qn("w:id"), str(comment_id))
    r_ref.append(ref_el)
    return r_ref


# ─── Cross-run helper: 複数runにまたがる find_str を扱う ─────────────────────

def _get_all_runs_info(para):
    """直接の <w:r> 子要素を (element, text, abs_start) のリストで返す。"""
    items = []
    abs_pos = 0
    for child in para._p:
        if child.tag == qn("w:r"):
            txt = "".join(t.text or "" for t in child.iter(qn("w:t")))
            items.append((child, txt, abs_pos))
            abs_pos += len(txt)
    return items


def _make_run_copy_with_text(run_el, text: str):
    """run_el を deepcopy してテキストだけを置き換えた要素を返す。"""
    new_run = copy.deepcopy(run_el)
    t_nodes = list(new_run.iter(qn("w:t")))
    t_el = t_nodes[0] if t_nodes else None
    for extra_t in t_nodes[1:]:
        parent = extra_t.getparent()
        if parent is not None:
            parent.remove(extra_t)
    if t_el is None:
        t_el = OxmlElement("w:t")
        new_run.append(t_el)
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return new_run


def _replace_text_cross_run(para, find_str: str, replace_str: str) -> bool:
    """
    非track版: find_str が複数 run にまたがる場合の置換。
    連結テキストで置換し、first run に結果を置き、残りの run を削除する。
    """
    r_items = _get_all_runs_info(para)
    if not r_items:
        return False
    full_text = "".join(txt for _, txt, _ in r_items)
    if find_str not in full_text:
        return False
    new_text = full_text.replace(find_str, replace_str)
    first_run_el = r_items[0][0]
    t_el = first_run_el.find(qn("w:t"))
    if t_el is None:
        t_el = OxmlElement("w:t")
        first_run_el.append(t_el)
    t_el.text = new_text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for run_el, _, _ in r_items[1:]:
        para._p.remove(run_el)
    return True


def _replace_text_cross_run_tracked(
    para, find_str: str, replace_str: str,
    author: str, date: str, rev: list,
    comments: "CommentsManager | None" = None,
) -> bool:
    """
    Track Changes版: find_str が複数 run にまたがる場合の置換。
    連結テキスト上で find_str を検索し、<w:del>/<w:ins> として挿入する。
    """
    r_items = _get_all_runs_info(para)
    if not r_items:
        return False
    full_text = "".join(txt for _, txt, _ in r_items)
    if find_str not in full_text:
        return False

    # 全出現箇所を収集
    occ_list = []
    sp = 0
    while True:
        idx = full_text.find(find_str, sp)
        if idx == -1:
            break
        occ_list.append((idx, idx + len(find_str)))
        sp = idx + 1

    if not occ_list:
        return False

    new_els = []
    cursor = 0
    r_idx = 0
    r_offset = 0

    def advance_plain(up_to):
        nonlocal cursor, r_idx, r_offset
        while cursor < up_to and r_idx < len(r_items):
            run_el, run_text, run_abs = r_items[r_idx]
            take_end = min(up_to, run_abs + len(run_text))
            take_text = run_text[r_offset: take_end - run_abs]
            if take_text:
                new_els.append(_make_run_copy_with_text(run_el, take_text))
            if take_end >= run_abs + len(run_text):
                r_idx += 1
                r_offset = 0
            else:
                r_offset = take_end - run_abs
            cursor = take_end

    def skip_to(up_to):
        nonlocal cursor, r_idx, r_offset
        while cursor < up_to and r_idx < len(r_items):
            run_el, run_text, run_abs = r_items[r_idx]
            skip_end = min(up_to, run_abs + len(run_text))
            if skip_end >= run_abs + len(run_text):
                r_idx += 1
                r_offset = 0
            else:
                r_offset = skip_end - run_abs
            cursor = skip_end

    for occ_start, occ_end in occ_list:
        advance_plain(occ_start)

        match_rPr = _rPr_copy(r_items[r_idx][0]) if r_idx < len(r_items) else OxmlElement("w:rPr")

        comment_id = None
        if comments is not None:
            comment_text = (
                f"「{find_str}」→「{replace_str}」" if replace_str
                else f"「{find_str}」を削除"
            )
            comment_id = comments.add_comment(comment_text, author, date)
            cs = OxmlElement("w:commentRangeStart")
            cs.set(qn("w:id"), str(comment_id))
            new_els.append(cs)

        new_els.append(_wrap_del(rev[0], author, date, _make_del_run(match_rPr, find_str)))
        rev[0] += 1

        if replace_str:
            new_els.append(_wrap_ins(rev[0], author, date, _make_ins_run(match_rPr, replace_str)))
            rev[0] += 1

        if comment_id is not None:
            ce = OxmlElement("w:commentRangeEnd")
            ce.set(qn("w:id"), str(comment_id))
            new_els.append(ce)
            new_els.append(_make_comment_ref_run(comment_id))

        skip_to(occ_end)

    advance_plain(len(full_text))

    # 先頭 run の直前に新要素を順番通りに挿入し、旧 run を全て削除
    first_run_el = r_items[0][0]
    for el in new_els:
        first_run_el.addprevious(el)
    for run_el, _, _ in r_items:
        para._p.remove(run_el)

    return True


def apply_replace_text_tracked(

    doc,
    replacements: list,
    author: str,
    date: str,
    rev: list,
    comments: "CommentsManager | None" = None,
) -> int:
    """
    replaceText を Track Changes として記録しながら適用する。
    comments が指定されている場合、各置換箇所に Word コメントを挿入する。
    """
    changed = 0
    for para in iter_all_paragraphs(doc):
        para_changed = False
        for rep in replacements:
            find_str = rep.get("find", "")
            replace_str = rep.get("replace", "")
            if not find_str:
                continue

            single_run_found = False
            for run in list(para.runs):
                if find_str not in run.text:
                    continue

                single_run_found = True
                original_text = run.text
                rPr_elem = _rPr_copy(run._r)
                parent = run._r.getparent()
                pos = list(parent).index(run._r)
                parent.remove(run._r)

                # コメント + commentRangeStart を挿入
                comment_id = None
                if comments is not None:
                    comment_text = (
                        f"「{find_str}」→「{replace_str}」"
                        if replace_str
                        else f"「{find_str}」を削除"
                    )
                    comment_id = comments.add_comment(comment_text, author, date)
                    cs = OxmlElement("w:commentRangeStart")
                    cs.set(qn("w:id"), str(comment_id))
                    parent.insert(pos, cs)
                    pos += 1

                # テキストを find_str で分割し、del/ins を挿入する
                parts = original_text.split(find_str)
                insert_pos = pos
                for i, part in enumerate(parts):
                    if i > 0:
                        del_el = _wrap_del(rev[0], author, date, _make_del_run(rPr_elem, find_str))
                        rev[0] += 1
                        parent.insert(insert_pos, del_el)
                        insert_pos += 1
                        if replace_str:
                            ins_el = _wrap_ins(rev[0], author, date, _make_ins_run(rPr_elem, replace_str))
                            rev[0] += 1
                            parent.insert(insert_pos, ins_el)
                            insert_pos += 1
                    if part:
                        parent.insert(insert_pos, _make_plain_run(rPr_elem, part))
                        insert_pos += 1

                # commentRangeEnd + commentReference を挿入
                if comment_id is not None:
                    ce = OxmlElement("w:commentRangeEnd")
                    ce.set(qn("w:id"), str(comment_id))
                    parent.insert(insert_pos, ce)
                    insert_pos += 1
                    parent.insert(insert_pos, _make_comment_ref_run(comment_id))

                para_changed = True

            # Cross-run fallback: find_str が複数 run にまたがる場合
            if not single_run_found and find_str in para.text:
                if _replace_text_cross_run_tracked(
                    para, find_str, replace_str, author, date, rev, comments
                ):
                    para_changed = True
        if para_changed:
            changed += 1
    return changed


def apply_format_runs_tracked(
    doc,
    formats: list,
    author: str,
    date: str,
    rev: list,
    comments: "CommentsManager | None" = None,
) -> int:
    """
    formatRuns を <w:rPrChange> を使った Track Changes として記録しながら適用する。
    comments が指定されている場合、変更段落の先頭にコメントを挿入する。
    """
    changed = 0
    for fmt in formats:
        match_text = str(fmt.get("matchText") or "").strip()
        apply_all = not match_text
        bold = fmt.get("bold")
        italic = fmt.get("italic")
        font_size = fmt.get("fontSize")
        font_color = parse_hex_color(fmt.get("fontColor"))
        font_face = str(fmt.get("fontFace") or "").strip() or None

        if bold is None and italic is None and font_size is None and font_color is None and font_face is None:
            continue

        # このフォーマット操作のコメント文を生成
        if comments is not None:
            fmt_parts = []
            if bold is not None:
                fmt_parts.append(f"太字={'ON' if bold else 'OFF'}")
            if italic is not None:
                fmt_parts.append(f"斜体={'ON' if italic else 'OFF'}")
            if font_size is not None:
                fmt_parts.append(f"フォントサイズ={font_size}pt")
            if font_color is not None:
                fmt_parts.append(f"文字色=#{fmt.get('fontColor', '')}")
            if font_face is not None:
                fmt_parts.append(f"フォント={font_face}")
            fmt_desc = "書式変更: " + "、".join(fmt_parts)

        for para in iter_all_paragraphs(doc):
            if not apply_all and match_text not in para.text:
                continue
            runs = para.runs
            if not runs:
                continue

            # コメント + commentRangeStart を先頭ランの前に挿入
            comment_id = None
            if comments is not None:
                comment_id = comments.add_comment(fmt_desc, author, date)
                first_run_el = runs[0]._r
                parent = first_run_el.getparent()
                first_run_pos = list(parent).index(first_run_el)
                cs = OxmlElement("w:commentRangeStart")
                cs.set(qn("w:id"), str(comment_id))
                parent.insert(first_run_pos, cs)
                # first_run_pos は +1 されたが runs[0]._r の参照はそのまま有効

            for run in runs:
                orig_rPr = copy.deepcopy(run._r.find(qn("w:rPr")) or OxmlElement("w:rPr"))

                if bold is not None:
                    run.bold = bold
                if italic is not None:
                    run.italic = italic
                if font_size is not None:
                    run.font.size = Pt(float(font_size))
                if font_color is not None:
                    run.font.color.rgb = font_color
                if font_face is not None:
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), font_face)
                    rFonts.set(qn("w:hAnsi"), font_face)
                    rFonts.set(qn("w:eastAsia"), font_face)
                    rFonts.set(qn("w:cs"), font_face)

                new_rPr = run._r.get_or_add_rPr()
                existing = new_rPr.find(qn("w:rPrChange"))
                if existing is not None:
                    new_rPr.remove(existing)
                rPrChange = OxmlElement("w:rPrChange")
                rPrChange.set(qn("w:id"), str(rev[0]))
                rPrChange.set(qn("w:author"), author)
                rPrChange.set(qn("w:date"), date)
                rev[0] += 1
                rPrChange.append(orig_rPr)
                new_rPr.append(rPrChange)

            # commentRangeEnd + commentReference を末尾ランの後に挿入
            if comment_id is not None:
                last_run_el = runs[-1]._r
                parent = last_run_el.getparent()
                last_run_pos = list(parent).index(last_run_el)
                ce = OxmlElement("w:commentRangeEnd")
                ce.set(qn("w:id"), str(comment_id))
                parent.insert(last_run_pos + 1, ce)
                parent.insert(last_run_pos + 2, _make_comment_ref_run(comment_id))

            changed += 1
    return changed


def apply_add_paragraphs_tracked(doc, paragraphs: list, author: str, date: str, rev: list) -> int:
    """
    addParagraphs を <w:ins> を使った Track Changes として記録しながら追加する。
    段落マークと段落テキストの両方に <w:ins> を付与する。
    """
    added = 0
    for para_def in paragraphs:
        text = str(para_def.get("text") or "").strip()
        if not text:
            continue
        style = str(para_def.get("style") or "Normal").strip()
        bold = para_def.get("bold")
        italic = para_def.get("italic")
        font_size = para_def.get("fontSize")
        font_color_hex = str(para_def.get("fontColor") or "").replace("#", "").strip().upper() or None
        font_face = str(para_def.get("fontFace") or "").strip() or None

        try:
            para = doc.add_paragraph("", style=style)
        except Exception:
            para = doc.add_paragraph("")

        # 段落マーク（¶）自体を挿入済みとしてマーク
        pPr = para._p.get_or_add_pPr()
        rPr_in_pPr = pPr.find(qn("w:rPr"))
        if rPr_in_pPr is None:
            rPr_in_pPr = OxmlElement("w:rPr")
            pPr.append(rPr_in_pPr)
        ins_mark = OxmlElement("w:ins")
        ins_mark.set(qn("w:id"), str(rev[0]))
        ins_mark.set(qn("w:author"), author)
        ins_mark.set(qn("w:date"), date)
        rev[0] += 1
        rPr_in_pPr.append(ins_mark)

        # テキストを持つ run を組み立てる
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        if font_size:
            half_pt = str(int(float(font_size) * 2))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), half_pt)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), half_pt)
            rPr.append(sz)
            rPr.append(szCs)
        if font_color_hex and len(font_color_hex) == 6:
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), font_color_hex)
            rPr.append(color_el)
        if font_face:
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), font_face)
            rFonts.set(qn("w:hAnsi"), font_face)
            rFonts.set(qn("w:eastAsia"), font_face)
            rFonts.set(qn("w:cs"), font_face)
            rPr.insert(0, rFonts)
        r.append(rPr)

        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)

        ins_el = _wrap_ins(rev[0], author, date, r)
        rev[0] += 1
        para._p.append(ins_el)
        added += 1

    return added


# ─── エントリポイント ────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--plan", required=True)
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    plan_path = Path(args.plan)

    plan: dict[str, Any] = json.loads(plan_path.read_text(encoding="utf-8"))
    doc = Document(str(input_path))
    total_paragraphs = sum(1 for _ in iter_all_paragraphs(doc))
    changed = 0

    track_changes: bool = bool(plan.get("trackChanges", False))

    if track_changes:
        date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        rev = [1]
        comments = CommentsManager(doc, rev)
        changed += apply_replace_text_tracked(
            doc, plan.get("replaceText") or [], TRACK_AUTHOR, date, rev, comments
        )
        changed += apply_format_runs_tracked(
            doc, plan.get("formatRuns") or [], TRACK_AUTHOR, date, rev, comments
        )
        changed += apply_add_paragraphs_tracked(
            doc, plan.get("addParagraphs") or [], TRACK_AUTHOR, date, rev
        )
        comments.flush()
    else:
        changed += apply_replace_text(doc, plan.get("replaceText") or [])
        changed += apply_format_runs(doc, plan.get("formatRuns") or [])
        changed += apply_add_paragraphs(doc, plan.get("addParagraphs") or [])

    doc.save(str(output_path))
    print(json.dumps({"changedParagraphs": changed, "totalParagraphs": total_paragraphs}))


if __name__ == "__main__":
    main()
